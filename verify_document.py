"""
Verify the created Word document content
"""

from docx import Document

def verify_document():
    """Read and display the contents of the created Word document"""

    doc_path = r'C:\myproject\output\reports\briefing_2025-12-27.docx'
    doc = Document(doc_path)

    print("Document Verification")
    print("=" * 60)
    print(f"File: {doc_path}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("=" * 60)
    print()

    print("DOCUMENT CONTENT:")
    print("-" * 60)

    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():  # Only print non-empty paragraphs
            # Show style information
            style_name = para.style.name
            text = para.text

            # Check for color formatting in runs
            has_color = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    has_color = True
                    rgb = run.font.color.rgb
                    print(f"{i}. [{style_name}] {text}")
                    print(f"   -> Contains colored text: RGB({rgb[0]}, {rgb[1]}, {rgb[2]})")
                    break

            if not has_color:
                print(f"{i}. [{style_name}] {text}")

    print("-" * 60)
    print("\nDocument created successfully with professional formatting!")

if __name__ == "__main__":
    verify_document()
