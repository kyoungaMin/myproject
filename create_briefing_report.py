"""
Stock Market Briefing Report Generator
Creates a professional Word document with formatted financial briefing content
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_briefing_report():
    """Create a professional stock market briefing report in Word format"""

    # Create a new Document
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('당신이 잠든 사이 | 2025.12.05', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    title_run.font.bold = True

    # Add spacing
    doc.add_paragraph()

    # Section 1: Today's Hot Stock
    heading1 = doc.add_heading('오늘의 화제 종목: TESLA (TSLA)', level=2)
    heading1_run = heading1.runs[0]
    heading1_run.font.size = Pt(16)
    heading1_run.font.color.rgb = RGBColor(0, 51, 102)

    # Stock details with bullet points
    stock_info = [
        ('주가: $385.20 ', '+8.7%', True),  # True indicates green color for percentage
        ('선정 기준: 거래량 1위', '', False)
    ]

    for item, highlight, is_positive in stock_info:
        p = doc.add_paragraph(style='List Bullet')
        run1 = p.add_run(item)
        run1.font.size = Pt(11)

        if highlight:
            run2 = p.add_run(highlight)
            run2.font.size = Pt(11)
            run2.font.bold = True
            if is_positive:
                run2.font.color.rgb = RGBColor(0, 176, 80)  # Green for positive percentage

    # Add spacing
    doc.add_paragraph()

    # Section 2: Why is it trending?
    heading2 = doc.add_heading('왜 화제인가?', level=2)
    heading2_run = heading2.runs[0]
    heading2_run.font.size = Pt(16)
    heading2_run.font.color.rgb = RGBColor(0, 51, 102)

    # Reasons with bullet points
    reasons = [
        '사이버트럭 판매량 급증',
        'FSD v13 출시 예고'
    ]

    for reason in reasons:
        p = doc.add_paragraph(reason, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    # Add spacing
    doc.add_paragraph()

    # Section 3: Related News TOP 3
    heading3 = doc.add_heading('관련 뉴스 TOP 3', level=2)
    heading3_run = heading3.runs[0]
    heading3_run.font.size = Pt(16)
    heading3_run.font.color.rgb = RGBColor(0, 51, 102)

    # News items with numbered list
    news_items = [
        '테슬라 사이버트럭 미국 판매 3위 - Reuters',
        'FSD v13 무감독 자율주행 - Bloomberg',
        '중국 시장 반등 - CNBC'
    ]

    for i, news in enumerate(news_items, 1):
        p = doc.add_paragraph(style='List Number')
        p.text = news
        p.runs[0].font.size = Pt(11)

    # Add footer space
    doc.add_paragraph()

    # Add a subtle footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run('Stock Market Briefing Report')
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

    # Save the document
    output_path = r'C:\myproject\output\reports\briefing_2025-12-27.docx'
    doc.save(output_path)

    print(f"Document successfully created: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        file_path = create_briefing_report()
        print(f"\n✓ Professional briefing report created successfully!")
        print(f"✓ Location: {file_path}")
        print(f"✓ Format: Microsoft Word (.docx)")
        print(f"✓ Formatting applied:")
        print(f"  - Professional heading styles")
        print(f"  - Green color for positive stock movement (+8.7%)")
        print(f"  - Bullet points for key information")
        print(f"  - Numbered list for news items")
        print(f"  - Corporate color scheme (dark blue headings)")
    except Exception as e:
        print(f"Error creating document: {e}")
        raise
